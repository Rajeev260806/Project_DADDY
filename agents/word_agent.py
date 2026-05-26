import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from loguru import logger
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage, SystemMessage
from agents.base_agent import BaseAgent
from config import OLLAMA_MODEL, OLLAMA_BASE_URL

SHORTCUTS = {
    "desktop":   str(Path.home() / "Desktop"),
    "documents": str(Path.home() / "Documents"),
    "downloads": str(Path.home() / "Downloads"),
    "home":      str(Path.home()),
}

WORD_AGENT_PROMPT = """You are a Microsoft Word command parser.
The user will give a plain English command about a Word document.
Respond ONLY with a JSON object — no explanation, no markdown, no extra text.

Supported actions:

1. Create a new Word document:
{"action":"create","path":"Desktop","name":"report.docx","title":"My Report"}

2. Read a document:
{"action":"read","path":"Desktop","name":"report.docx"}

3. Summarize a document:
{"action":"summarize","path":"Desktop","name":"report.docx"}

4. Add a paragraph:
{"action":"add_paragraph","path":"Desktop","name":"report.docx","text":"This is a new paragraph.","style":"Normal"}

5. Add a heading:
{"action":"add_heading","path":"Desktop","name":"report.docx","text":"Introduction","level":1}

6. Replace text:
{"action":"replace_text","path":"Desktop","name":"report.docx","find":"old text","replace":"new text"}

7. Add a table:
{"action":"add_table","path":"Desktop","name":"report.docx","headers":["Name","Age","City"],"rows":[["Arun","25","Kerala"]]}

8. Word count:
{"action":"word_count","path":"Desktop","name":"report.docx"}

9. Delete a paragraph containing specific text:
{"action":"delete_paragraph","path":"Desktop","name":"report.docx","contains":"text to find"}

10. Set font for entire document:
{"action":"set_font","path":"Desktop","name":"report.docx","font_name":"Arial","font_size":12}

Rules:
- path uses short names: Desktop, Documents, Downloads
- name always ends with .docx
- style options: Normal, Heading1, Heading2, Quote, ListBullet
- heading level 1 is biggest, level 3 is smallest
- Respond ONLY with JSON. Nothing else.
"""

class WordAgent(BaseAgent):
    
    def __init__(self):
        super().__init__("WordAgent")
        self.llm = ChatOllama(
            model = OLLAMA_MODEL,
            base_url= OLLAMA_BASE_URL,
            temperature=0
        )

    def resolve_path(self, folder: str) -> Path:
        folder = folder.strip()
        for alias, full_path in SHORTCUTS.items():
            if folder.lower().startswith(alias):
                remainder = folder[len(alias):].lstrip("/\\")
                if remainder:
                    return Path(full_path) / remainder
                return Path(full_path)
        p = Path(folder)
        if p.is_absolute():
            return p
        return Path.home() / folder
    
    def parse_command(self, command: str) -> dict:
        messages = [
            SystemMessage(content=WORD_AGENT_PROMPT),
            HumanMessage(content=command)
        ]
        response = self.llm.invoke(messages)
        raw = response.content.strip()
        if raw.startswith("```"):
            lines = raw.split("\n")
            raw = "\n".join(lines[1:-1])
        return json.loads(raw)
    
    def get_doc_path(self,data:dict)->Path:
        base = self.resolve_path(data.get("path","documents"))
        return Path(base) / data["name"]
    
    def load_doc(self,data:dict)->tuple[Document,Path]:
        file_path = self.get_doc_path(data)
        if not file_path.exists:
            raise FileNotFoundError(f"'{data['name']}' not found in "
                f"{self.resolve_path(data.get('path', 'documents')).name}.")
        return Document(file_path),file_path
    
    def extract_text(self, doc: Document) -> str:
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    def create(self, data: dict) -> str:
        base = self.resolve_path(data.get("path", "documents"))
        doc_path = base / data["name"]
        base.mkdir(parents=True, exist_ok=True)
        doc = Document()
        title = data.get("title", data["name"].replace(".docx", ""))
        doc.add_heading(title, level=0)
        doc.add_paragraph("")
        doc.save(str(doc_path))
        logger.success(f"Created Word document: {doc_path}")
        return f"Done! '{data['name']}' created in {base.name}."
    
    def read(self,data:dict)->str:
        doc, _ = self.load_doc(data)
        text   = self.extract_text(doc)
        if not text:
            return f"'{data['name']}' exists but has no readable text."
        preview = text[:1500]
        if len(text) > 1500:
            preview += f"\n... (showing first 1500 of {len(text)} characters)"
        return f"Contents of '{data['name']}':\n{preview}"
    
    def summarize(self, data: dict) -> str:
        doc, _ = self.load_doc(data)
        text = self.extract_text(doc)
        if not text:
            return f"'{data['name']}' has no text to summarize."
        prompt = f"Summarize this document clearly in 5 sentences or less:\n\n{text[:4000]}"
        response = self.llm.invoke([HumanMessage(content=prompt)])
        return f"Summary of '{data['name']}':\n{response.content}"
    
    def add_paragraph(self, data: dict) -> str:
        doc, doc_path = self.load_doc(data)
        style = data.get("style", "Normal")
        try:
            doc.add_paragraph(data["text"], style=style)
        except Exception:
            doc.add_paragraph(data["text"])
        doc.save(str(doc_path))
        return f"Done! Paragraph added to '{data['name']}'."
    
    def add_heading(self, data: dict) -> str:
        doc, doc_path = self.load_doc(data)
        doc.add_heading(data["text"], level=int(data.get("level", 1)))
        doc.save(str(doc_path))
        return f"Done! Heading '{data['text']}' added to '{data['name']}'."
    
    def replace_text(self, data: dict) -> str:
        doc, doc_path = self.load_doc(data)
        find    = data["find"]
        replace = data["replace"]
        count   = 0
        for para in doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find in para.text:
                            for run in para.runs:
                                if find in run.text:
                                    run.text = run.text.replace(find, replace)
                                    count += 1
        if count == 0:
            return f"Text '{find}' not found in '{data['name']}'."
        doc.save(str(doc_path))
        return f"Done! Replaced '{find}' with '{replace}' — {count} occurrence(s) in '{data['name']}'."

    def add_table(self, data: dict) -> str:
        doc, doc_path = self.load_doc(data)
        headers = data.get("headers", [])
        rows    = data.get("rows", [])
        if not headers:
            return "No headers provided for the table."
        table  = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row_data in rows:
            row = table.add_row()
            for i, val in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = str(val)
        doc.add_paragraph("")
        doc.save(str(doc_path))
        return f"Done! Table with {len(rows)} row(s) added to '{data['name']}'."
    
    def word_count(self, data: dict) -> str:
        doc, _  = self.load_doc(data)
        text    = self.extract_text(doc)
        words   = len(text.split())
        chars   = len(text)
        paras   = len([p for p in doc.paragraphs if p.text.strip()])
        return (
            f"'{data['name']}' statistics:\n"
            f"  Words      : {words}\n"
            f"  Characters : {chars}\n"
            f"  Paragraphs : {paras}"
        )

    def delete_paragraph(self, data: dict) -> str:
        doc, doc_path = self.load_doc(data)
        contains = data.get("contains", "")
        count    = 0
        for para in doc.paragraphs:
            if contains.lower() in para.text.lower():
                p = para._element
                p.getparent().remove(p)
                count += 1
        if count == 0:
            return f"No paragraph containing '{contains}' found in '{data['name']}'."
        doc.save(str(doc_path))
        return f"Done! {count} paragraph(s) containing '{contains}' deleted from '{data['name']}'."

    def set_font(self, data: dict) -> str:
        doc, doc_path = self.load_doc(data)
        font_name = data.get("font_name", "Calibri")
        font_size = int(data.get("font_size", 11))
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
        doc.save(str(doc_path))
        return f"Done! Font set to {font_name} {font_size}pt in '{data['name']}'."
    
    def handle(self, command: str) -> str:
        logger.info(f"WordAgent handling: {command}")
        data   = self.parse_command(command)
        action = data.get("action", "")
        logger.debug(f"Parsed action: {data}")

        actions = {
            "create":           self.create,
            "read":             self.read,
            "summarize":        self.summarize,
            "add_paragraph":    self.add_paragraph,
            "add_heading":      self.add_heading,
            "replace_text":     self.replace_text,
            "add_table":        self.add_table,
            "word_count":       self.word_count,
            "delete_paragraph": self.delete_paragraph,
            "set_font":         self.set_font,
        }

        if action not in actions:
            return f"Sorry, I don't know how to do '{action}' on a Word document."

        return actions[action](data)